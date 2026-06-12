"""Document parsers: bytes in, plain text out.

Format detection prefers the filename extension, falling back to the declared
content type. Parsers are synchronous and CPU-bound — callers run them in a
worker thread. Extraction is deliberately plain-text-first: layout-aware
extraction (tables, reading order) is a later refinement once the eval
pipeline can measure whether it helps retrieval.
"""

import io

import docx
import pymupdf
from bs4 import BeautifulSoup


class UnsupportedDocumentType(Exception):
    def __init__(self, detected: str):
        super().__init__(f"Unsupported document type: {detected}")
        self.detected = detected


class DocumentParseError(Exception):
    pass


_EXTENSION_KINDS = {
    "pdf": "pdf",
    "docx": "docx",
    "html": "html",
    "htm": "html",
    "txt": "text",
    "md": "text",
    "markdown": "text",
}

_CONTENT_TYPE_KINDS = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/html": "html",
    "text/plain": "text",
    "text/markdown": "text",
}


def _detect(filename: str, content_type: str | None) -> str:
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    kind = _EXTENSION_KINDS.get(extension)
    if kind is None and content_type:
        kind = _CONTENT_TYPE_KINDS.get(content_type.split(";")[0].strip().lower())
    if kind is None:
        raise UnsupportedDocumentType(extension or content_type or "unknown")
    return kind


def _parse_pdf(data: bytes) -> str:
    with pymupdf.open(stream=data, filetype="pdf") as document:
        return "\n".join(page.get_text() for page in document)


def _parse_docx(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(part for part in parts if part.strip())


def _parse_html(data: bytes) -> str:
    soup = BeautifulSoup(data, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(" ", strip=True)


_PARSERS = {
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "html": _parse_html,
    "text": lambda data: data.decode("utf-8", errors="replace"),
}


def parse_document(data: bytes, filename: str, content_type: str | None) -> str:
    kind = _detect(filename, content_type)
    try:
        return _PARSERS[kind](data)
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse {kind} document: {exc}") from exc
