"""Document parser + upload endpoint tests. Fixture files are generated
in-test with the same libraries (synthetic PII only, per TESTING.md)."""

import io

import docx
import pymupdf
import pytest

from forge.rag.parsing import (
    DocumentParseError,
    UnsupportedDocumentType,
    parse_document,
)

NOTE = "Patient John Smith (SSN 536-90-4399) is prescribed Metformin 1000mg daily."


def _pdf_bytes(text: str) -> bytes:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    return document.tobytes()


def _docx_bytes(text: str) -> bytes:
    document = docx.Document()
    document.add_paragraph(text)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Medication"
    table.rows[0].cells[1].text = "Metformin"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


HTML = b"""<html><head><style>body{color:red}</style>
<script>alert('ignore me')</script></head>
<body><h1>Policy</h1><p>The refund window is 47 days.</p></body></html>"""


# --- parser unit tests --------------------------------------------------------


def test_pdf_roundtrip():
    assert "536-90-4399" in parse_document(_pdf_bytes(NOTE), "note.pdf", None)


def test_docx_includes_paragraphs_and_tables():
    text = parse_document(_docx_bytes(NOTE), "note.docx", None)
    assert "John Smith" in text
    assert "Medication | Metformin" in text


def test_html_strips_script_and_style():
    text = parse_document(HTML, "policy.html", None)
    assert "refund window is 47 days" in text
    assert "alert" not in text
    assert "color:red" not in text


def test_detection_falls_back_to_content_type():
    assert "hello" in parse_document(b"hello", "no-extension", "text/plain")


def test_unsupported_type_and_corrupt_file():
    with pytest.raises(UnsupportedDocumentType):
        parse_document(b"\x89PNG", "image.png", "image/png")
    with pytest.raises(DocumentParseError):
        parse_document(b"not a real pdf", "broken.pdf", None)


# --- upload endpoint ----------------------------------------------------------


async def test_pdf_upload_is_parsed_scrubbed_and_ingested(
    client, auth_headers, fake_embeddings
):
    response = await client.post(
        "/v1/documents/upload",
        headers=auth_headers,
        files={"file": ("note.pdf", _pdf_bytes(NOTE), "application/pdf")},
    )
    assert response.status_code == 201
    body = response.json()
    assert body["chunks"] == 1
    assert body["pii_redactions"] >= 2  # parser path feeds the same PII pipeline


async def test_upload_title_defaults_to_filename(client, auth_headers, fake_embeddings):
    response = await client.post(
        "/v1/documents/upload",
        headers=auth_headers,
        files={"file": ("policy.html", HTML, "text/html")},
    )
    assert response.status_code == 201

    search = await client.post(
        "/v1/search", headers=auth_headers, json={"query": "return period", "limit": 1}
    )
    (result,) = search.json()["data"]
    assert result["title"] == "policy.html"
    assert "47 days" in result["text"]


async def test_upload_rejects_unsupported_and_empty(client, auth_headers):
    response = await client.post(
        "/v1/documents/upload",
        headers=auth_headers,
        files={"file": ("img.png", b"\x89PNG", "image/png")},
    )
    assert response.status_code == 415

    empty_pdf = pymupdf.open()
    empty_pdf.new_page()
    response = await client.post(
        "/v1/documents/upload",
        headers=auth_headers,
        files={"file": ("empty.pdf", empty_pdf.tobytes(), "application/pdf")},
    )
    assert response.status_code == 422


async def test_upload_size_cap(client, auth_headers, monkeypatch):
    from forge.config import get_settings

    monkeypatch.setattr(get_settings(), "rag_max_upload_bytes", 100)
    response = await client.post(
        "/v1/documents/upload",
        headers=auth_headers,
        files={"file": ("big.txt", b"x" * 200, "text/plain")},
    )
    assert response.status_code == 413